#!/usr/bin/env python3
"""Build polished Elsevier revision Word uploads (especially blind manuscript)."""
from __future__ import annotations

import re
import shutil
import subprocess
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
OUT = PAPER / "em_revision"
FIG = PAPER / "figures"
REF = OUT / "reference_sat.docx"

MANUSCRIPT_TITLE = (
    "Task-Oriented Semantic-Guided Restoration for Agricultural UAV Crop/Weed "
    "Segmentation under Structure and Composite Degradation"
)

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
TABLE_SIZE = Pt(10)
CAPTION_SIZE = Pt(10)
TITLE_SIZE = Pt(14)
H1_SIZE = Pt(14)
H2_SIZE = Pt(12)
H3_SIZE = Pt(12)
LINE_SP = 1.15

FIGURE_FILES = [
    "fig2_framework.png",
    "fig2_example.png",
    "fig3_baselines_weeds.png",
    "fig4_when_helps.png",
    "guide_quality_curve.png",
    "fig4_boundary_tradeoff.png",
    "fig5_sobel_sensitivity.png",
    "fields_example.png",
]

# All manuscript tables (caption, header row, body rows)
MANUSCRIPT_TABLES: list[tuple[str, list[list[str]]]] = [
    (
        "Table 1. Datasets, degradation settings, and evaluation roles. n denotes held-out "
        "evaluation crops per run (256×256; paired over seeds and degradation configs).",
        [
            ["Dataset", "Domain", "Task", "n", "Degradation", "Metrics", "Purpose"],
            ["WeedsGalore", "UAV maize", "crop/weed", "468", "struct./veil/comp.", "PSNR, SSIM, mIoU", "Primary in-domain UAV crop/weed"],
            ["CWFID", "ground carrot", "plant/soil", "320", "struct./comp.", "mIoU", "Agricultural replication"],
            ["CoFly", "real UAV cotton", "weed/rest", "320", "struct./comp.", "mIoU", "Real-UAV replication"],
            ["LoveDA-R", "RS rural (OOD)", "land cover", "160", "struct./comp.", "PSNR, SSIM, mIoU", "OOD scope test (appendix)"],
            ["LoveDA-U", "RS urban (OOD)", "land cover", "160", "struct./comp.", "PSNR, SSIM, mIoU", "OOD scope test (appendix)"],
        ],
    ),
    (
        "Table 2. Main task-oriented results on agricultural datasets (frozen SegFormer-B0). "
        "Regime sub-rows report ΔmIoU vs bicubic.",
        [
            ["Dataset", "Regime", "Bicubic mIoU", "Ours mIoU", "ΔmIoU", "p", "ΔPSNR", "ΔSSIM"],
            ["WeedsGalore", "all", "0.390", "0.453", "+0.063", "<10⁻¹³", "+4.46", "+0.033"],
            ["WeedsGalore", "structure", "—", "—", "+0.102", "<10⁻¹³", "—", "—"],
            ["WeedsGalore", "composite", "—", "—", "+0.119", "<10⁻¹³", "—", "—"],
            ["WeedsGalore", "veil", "—", "—", "−0.032", "<10⁻¹³", "—", "—"],
            ["CWFID", "s+c", "0.661", "0.714", "+0.053", "6×10⁻¹⁷", "+2.08", "+0.030"],
            ["CWFID", "structure", "—", "—", "+0.021", "—", "—", "—"],
            ["CWFID", "composite", "—", "—", "+0.085", "—", "—", "—"],
            ["CoFly", "s+c", "0.353", "0.412", "+0.058", "2×10⁻¹²", "+0.67", "+0.009"],
            ["CoFly", "structure", "—", "—", "+0.026", "—", "—", "—"],
            ["CoFly", "composite", "—", "—", "+0.090", "—", "—", "—"],
        ],
    ),
    (
        "Table 3. General-purpose SR baselines under the no-leak protocol (WeedsGalore, "
        "frozen SegFormer-B0, ×4, 2 seeds). Lower BRISQUE is better.",
        [
            ["Group", "Method", "mIoU", "ΔmIoU", "BRISQUE", "Conclusion"],
            ["Reference", "Bicubic", "0.609", "—", "56.7", "baseline"],
            ["Reference", "Tiny ResCNN", "0.615", "+0.006", "55.3", "helpful"],
            ["Transformer SR", "SwinIR", "0.589", "−0.021", "47.1", "visually better, task worse"],
            ["Transformer SR", "HAT-L", "0.590", "−0.019", "46.6", "best BRISQUE, task worse"],
            ["Transformer SR", "Restormer", "0.606", "−0.003", "56.7", "neutral"],
            ["Transformer SR", "Swin2SR (classical)", "0.593", "−0.017", "47.9", "visually better, task worse"],
            ["Transformer SR", "Swin2SR (real-world)", "0.573", "−0.036", "63.8", "harmful"],
            ["Implicit", "LIIF", "0.583", "−0.026", "48.5", "visually better, task worse"],
            ["Ours", "Semantic-INR", "0.610", "+0.001 (n.s.)", "57.5", "task-neutral at low damage"],
        ],
    ),
    (
        "Table 4. Ablation on WeedsGalore (frozen SegFormer-B0). Δ vs bicubic.",
        [
            ["Variant", "ΔmIoU", "ΔPSNR", "ΔSSIM", "Interpretation"],
            ["Semantic-INR (full)", "+0.114", "+2.20", "+0.016", "task + fidelity balance"],
            ["− semantic", "+0.084", "+2.18", "+0.016", "task driver removed"],
            ["− texture", "+0.124", "+2.21", "+0.017", "fidelity branch; mIoU not primary"],
            ["− frequency loss", "+0.120", "+2.21", "+0.017", "fidelity branch; mIoU not primary"],
        ],
    ),
    (
        "Table A1. Out-of-domain scope: visual quality vs frozen downstream mIoU (SegFormer-B0; LoveDA OOD).",
        [
            ["Dataset", "Method", "PSNR", "SSIM", "BF", "mIoU"],
            ["LoveDA-R (OOD)", "Bicubic", "20.8", "0.454", "0.464", "0.318"],
            ["LoveDA-R (OOD)", "Ours", "25.4", "0.523", "0.471", "0.325"],
            ["LoveDA-U (OOD)", "Bicubic", "20.1", "0.430", "0.491", "0.332"],
            ["LoveDA-U (OOD)", "Ours", "23.9", "0.469", "0.514", "0.327"],
        ],
    ),
    (
        "Table A2. WeedsGalore downstream mIoU Δ vs bicubic (3 seeds, DeepLabV3+ frozen backbone).",
        [
            ["Regime", "DeepLabV3+ ΔmIoU"],
            ["structure", "+0.044"],
            ["composite", "+0.032"],
            ["veil/fog", "+0.013"],
            ["overall", "+0.030"],
        ],
    ),
    (
        "Table A3. Full results matrix (LoveDA OOD scope and WeedsGalore baselines; frozen SegFormer-B0). "
        "Regime sub-rows: ΔmIoU only.",
        [
            ["Data", "Reg.", "Method", "PSNR", "SSIM", "BF", "mIoU", "ΔmIoU"],
            ["WeedsGalore", "all", "Bicubic", "16.523", "0.248", "0.429", "0.390", "—"],
            ["WeedsGalore", "all", "Ours", "20.986", "0.281", "0.424", "0.453", "+0.063"],
            ["WeedsGalore", "struct.", "Δ vs bic.", "—", "—", "—", "—", "+0.102"],
            ["WeedsGalore", "comp.", "Δ vs bic.", "—", "—", "—", "—", "+0.119"],
            ["WeedsGalore", "veil", "Δ vs bic.", "—", "—", "—", "—", "−0.032"],
            ["CWFID", "s+c", "Bicubic", "20.156", "0.505", "—", "0.661", "—"],
            ["CWFID", "s+c", "Ours", "22.239", "0.535", "—", "0.714", "+0.053"],
            ["CoFly", "s+c", "Bicubic", "17.390", "0.285", "—", "0.353", "—"],
            ["CoFly", "s+c", "Ours", "18.056", "0.294", "—", "0.412", "+0.058"],
            ["LoveDA-R", "s+c", "Bicubic", "20.754", "0.454", "0.464", "0.318", "—"],
            ["LoveDA-R", "s+c", "Ours", "25.399", "0.523", "0.471", "0.325", "+0.008"],
            ["LoveDA-U", "s+c", "Bicubic", "20.114", "0.430", "0.491", "0.332", "—"],
            ["LoveDA-U", "s+c", "Ours", "23.861", "0.469", "0.514", "0.327", "−0.005"],
        ],
    ),
]

# (needle in paragraph text, kind, table_index OR figure filename, caption for figures)
CONTENT_BLOCKS: list[tuple[str, str, str | int, str | None]] = [
    ("summarizes datasets and degradations", "table", 0, None),
    (
        "Figure 1 sketches the framework",
        "figure",
        "fig2_framework.png",
        "Figure 1. SG-FR framework overview. Prior learning (top): clean-trained semantic guide "
        "extraction and blind frequency-guided prior estimation. Restoration (middle): structure "
        "recovery, texture modulation, and semantic refinement; semantic branch auxiliary at training, "
        "inert at inference. Evaluation (bottom): frozen clean-trained segmenter without GT masks.",
    ),
    (
        "shows a WeedsGalore composite",
        "figure",
        "fig2_example.png",
        "Figure 2. Representative WeedsGalore composite example (×4). Top row: clean RGB, degraded "
        "input, restored output; bottom row: GT mask and frozen SegFormer-B0 predictions with "
        "zoom-in on crop/weed boundaries (red boxes).",
    ),
    ("reports agricultural main results", "table", 1, None),
    ("under the identical no-leak protocol", "table", 2, None),
    (
        "segmentation error maps",
        "figure",
        "fig3_baselines_weeds.png",
        "Figure 3. Qualitative comparison on WeedsGalore composite (×4): restored RGB, segmentation, "
        "and error map; strong SR baselines look sharper but increase segmentation errors.",
    ),
    (
        "Figure 4 consolidates three analyses",
        "figure",
        "fig4_when_helps.png",
        "Figure 4. When restoration helps agricultural perception. (A) Task damage vs task gain. "
        "(B) Boundary-F vs task gain trade-off. (C) Structure vs veil degradation response.",
    ),
    ("ablates components with dual metrics", "table", 3, None),
    (
        "Ground-truth masks are thus not required",
        "figure",
        "guide_quality_curve.png",
        "Figure 5. Downstream gain vs semantic-guide quality (deployment without GT masks).",
    ),
    ("reports LoveDA Rural/Urban results", "table", 4, None),
    ("gives the complete agricultural and baseline matrix", "table", 6, None),
    ("WeedsGalore downstream mIoU", "table", 5, None),
]

APPENDIX_FIGURES = [
    (
        "Individual panels of Fig",
        [
            (
                "fig4_boundary_tradeoff.png",
                "Figure A1. Boundary fidelity vs semantic utility (detail for Figure 4 panel B).",
            ),
            (
                "fig5_sobel_sensitivity.png",
                "Figure A2. Sobel-quantile sensitivity of Boundary-F Δ vs bicubic.",
            ),
            (
                "fields_example.png",
                "Figure A3. Interpretable fields: structure, texture, α, and semantic-probability fields.",
            ),
        ],
    ),
]

UNNUMBERED_HEADINGS = {
    "Declaration of competing interest",
    "Funding",
    "Data availability",
    "Reproducibility",
    "Abstract",
}


def write_blind_tex() -> Path:
    src = (PAPER / "main.tex").read_text(encoding="utf-8")
    src = re.sub(
        r"\\author\[1\]\{Minghao Li\\corref\{cor1\}\}\n"
        r"\\ead\{061123071@stu\.aqnu\.edu\.cn\}\n"
        r"\\cortext\[cor1\]\{Corresponding author\.\}\n"
        r"\\affiliation\[1\]\{organization=\{Anqing Normal University\}, city=\{Anqing\}, country=\{China\}\}\n\n",
        "",
        src,
    )
    src = re.sub(
        r"\\section\*\{CRediT authorship contribution statement\}\n"
        r"Minghao Li: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation,\n"
        r"Data curation, Writing---original draft, Writing---review \\& editing, Visualization, Project\n"
        r"administration\.\n\n",
        "",
        src,
    )
    src = src.replace(
        "available from the corresponding author upon reasonable request.",
        "available from the authors upon reasonable request.",
    )
    blind = PAPER / "main_blind.tex"
    blind.write_text(src, encoding="utf-8")
    return blind


def write_pandoc_tex(blind_tex: Path) -> Path:
    """Simplify LaTeX wrappers so pandoc converts more tables and text cleanly."""
    src = blind_tex.read_text(encoding="utf-8")
    inject = r"""
% Pandoc-friendly table wrappers
\renewcommand{\papertable}[1]{#1}
\renewcommand{\widepapertable}[1]{#1}
"""
    src = src.replace("\\begin{document}", inject + "\\begin{document}", 1)
    out = PAPER / "main_blind_pandoc.tex"
    out.write_text(src, encoding="utf-8")
    return out


def build_reference_doc() -> None:
    if not REF.exists():
        subprocess.run(
            ["pandoc", "--print-default-data-file", "reference.docx"],
            stdout=open(REF, "wb"),
            check=True,
        )
    doc = Document(str(REF))
    for name, size, bold in [
        ("Normal", BODY_SIZE, False),
        ("Body Text", BODY_SIZE, False),
        ("First Paragraph", BODY_SIZE, False),
        ("Abstract", BODY_SIZE, False),
        ("Compact", TABLE_SIZE, False),
        ("Caption", CAPTION_SIZE, False),
        ("Table Caption", CAPTION_SIZE, False),
        ("Image Caption", CAPTION_SIZE, False),
        ("Bibliography", Pt(11), False),
    ]:
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = FONT
            st.font.size = size
            st.font.bold = bold
            st.font.color.rgb = RGBColor(0, 0, 0)
            pf = st.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = LINE_SP
            pf.space_after = Pt(6)
    for name, size in [("Heading 1", H1_SIZE), ("Heading 2", H2_SIZE), ("Heading 3", H3_SIZE), ("Title", TITLE_SIZE)]:
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = FONT
            st.font.size = size
            st.font.bold = True
            st.font.color.rgb = RGBColor(0, 0, 0)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    _clear_document_body(doc)
    doc.save(str(REF))


def _clear_document_body(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(doc.tables):
        table._element.getparent().remove(table._element)


def _new_styled_document() -> Document:
    build_reference_doc()
    return Document(str(REF))


def _read_abstract_text() -> str:
    abs_lines = (PAPER / "abstract_sat.txt").read_text(encoding="utf-8").splitlines()
    abs_text = ""
    for line in abs_lines:
        if line.startswith("Agricultural UAV"):
            abs_text = line.strip()
        elif abs_text and line.strip() and not line.startswith("Word count"):
            abs_text += " " + line.strip()
    return _replace_unicode_dashes(abs_text)


def pandoc_to_docx(tex: Path, docx: Path) -> None:
    subprocess.run(
        [
            "pandoc",
            str(tex.name),
            "-o",
            str(docx),
            f"--resource-path={PAPER}:{FIG}",
            f"--reference-doc={REF}",
            "--bibliography=refs.bib",
            "--citeproc",
            "--wrap=none",
            "--mathml",
        ],
        cwd=PAPER,
        check=True,
    )


def _set_run_font(run, size=BODY_SIZE, bold=False, italic=False) -> None:
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rPr.insert(0, rFonts)


def _format_paragraph(paragraph, size=BODY_SIZE, bold=False, italic=False, align=None, space_before=0, space_after=6) -> None:
    if align is not None:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SP
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    for run in paragraph.runs:
        _set_run_font(run, size=size, bold=bold, italic=italic)


def _shade_cell(cell, fill: str | None) -> None:
    if not fill:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _replace_unicode_dashes(text: str) -> str:
    text = text.replace("\u2014", "-").replace("\u2013", "-").replace("\u2212", "-")
    text = re.sub(r"\s+,", ",", text)
    text = re.sub(r",\s*", ", ", text)
    return text


def _insert_block_after(paragraph, elements: list) -> None:
    anchor = paragraph._p
    for el in elements:
        anchor.addnext(el)
        anchor = el


def _strip_pandoc_figures(doc: Document) -> None:
    """Remove pandoc inline figures and their caption paragraphs; we re-insert at anchors."""
    remove: list = []
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if not p._element.xpath(".//wp:inline"):
            continue
        remove.append(p._element)
        for j in range(i + 1, min(i + 3, len(paras))):
            nxt = paras[j]
            t = nxt.text.strip()
            if (
                nxt.style.name in ("Image Caption", "Caption")
                or t.startswith("Figure")
                or t.startswith("Fig.")
                or "framework overview" in t
                or "Representative WeedsGalore" in t
                or "Qualitative comparison" in t
                or "When restoration helps" in t
                or "Downstream gain" in t
                or "Boundary fidelity" in t
                or "Sobel-quantile" in t
                or "Interpretable fields" in t
            ):
                remove.append(nxt._element)
                break
    for el in remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _insert_figure_after(doc: Document, paragraph, filename: str, caption: str) -> None:
    img = FIG / filename
    if not img.exists():
        img = PAPER / filename
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_before = Pt(8)
    pic.paragraph_format.space_after = Pt(4)
    pic.add_run().add_picture(str(img), width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    _format_paragraph(
        cap,
        size=CAPTION_SIZE,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    _insert_block_after(paragraph, [pic._p, cap._p])


def _norm(text: str) -> str:
    return text.replace("\xa0", " ").replace("\u2013", "-").replace("\u2014", "-")


def _insert_content_blocks(doc: Document) -> None:
    """Insert tables/figures after anchor paragraphs (bottom-up to preserve order)."""
    block_order = {b[0]: i for i, b in enumerate(CONTENT_BLOCKS)}
    matches: list[tuple[int, int, object, tuple]] = []

    for bi, block in enumerate(CONTENT_BLOCKS):
        needle, kind, payload, cap = block
        for i, p in enumerate(doc.paragraphs):
            if needle in _norm(p.text):
                matches.append((i, bi, p, block))
                break

    for figs in APPENDIX_FIGURES:
        _heading, fig_list = figs
        for i, p in enumerate(doc.paragraphs):
            if _heading in _norm(p.text):
                matches.append((i, 1000, p, ("__appendix__", "appendix_figs", fig_list, None)))
                break

    # Higher paragraph index first; within same paragraph, higher block index first
    # so earlier blocks in CONTENT_BLOCKS end up closer to the anchor paragraph.
    matches.sort(key=lambda x: (x[0], x[1]), reverse=True)

    for _, _, p, block in matches:
        needle, kind, payload, cap = block
        if kind == "table":
            caption, rows = MANUSCRIPT_TABLES[int(payload)]
            elems = _make_table_element(doc, caption, rows)
            _insert_block_after(p, elems)
        elif kind == "appendix_figs":
            elems = []
            for fname, caption in payload:
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.paragraph_format.space_before = Pt(8)
                pic.add_run().add_picture(str(FIG / fname), width=Inches(6.0))
                cp = doc.add_paragraph(caption)
                _format_paragraph(cp, size=CAPTION_SIZE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                elems.extend([pic._p, cp._p])
            _insert_block_after(p, elems)
        elif kind == "figure":
            _insert_figure_after(doc, p, str(payload), str(cap))


def _insert_table_after_caption(doc: Document, caption_prefix: str, table_index: int) -> None:
    """Insert an extra table immediately after an already-inserted table caption."""
    for i, p in enumerate(doc.paragraphs):
        if not p.text.startswith(caption_prefix):
            continue
        # Walk forward to the spacer paragraph after the table grid
        for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip() == "" and j > i + 1:
                caption, rows = MANUSCRIPT_TABLES[table_index]
                elems = _make_table_element(doc, caption, rows)
                _insert_block_after(doc.paragraphs[j], elems)
                return
        caption, rows = MANUSCRIPT_TABLES[table_index]
        elems = _make_table_element(doc, caption, rows)
        _insert_block_after(p, elems)
        return


def _number_headings(doc: Document) -> None:
    from docx.enum.text import WD_BREAK

    sec = sub = subsub = 0
    first_h1 = True
    for p in doc.paragraphs:
        raw = p.text.strip()
        if not raw:
            continue
        if p.style.name == "Heading 1":
            if any(raw.startswith(u) or u in raw for u in UNNUMBERED_HEADINGS):
                continue
            if not first_h1:
                run = p.runs[0] if p.runs else p.add_run()
                run.add_break(WD_BREAK.PAGE)
            first_h1 = False
            sec += 1
            sub = subsub = 0
            p.text = f"{sec}. {raw}"
            _format_paragraph(p, size=H1_SIZE, bold=True, space_before=12, space_after=6)
        elif p.style.name == "Heading 2":
            sub += 1
            subsub = 0
            label = raw.split(". ", 1)[-1] if re.match(r"^\d+\.\d+\s", raw) else raw
            p.text = f"{sec}.{sub} {label}"
            _format_paragraph(p, size=H2_SIZE, bold=True, space_before=10, space_after=4)
        elif p.style.name == "Heading 3":
            subsub += 1
            label = re.sub(r"^[\d.]+\s*", "", raw)
            p.text = f"{sec}.{sub}.{subsub} {label}"
            _format_paragraph(p, size=H3_SIZE, bold=True, space_before=8, space_after=4)


def _make_table_element(doc: Document, caption: str, rows: list[list[str]]):
    """Build caption paragraph + table; return XML elements to insert."""
    cap = doc.add_paragraph(caption)
    cap.paragraph_format.space_before = Pt(12)
    cap.paragraph_format.space_after = Pt(6)
    _format_paragraph(cap, size=CAPTION_SIZE, bold=True)

    nrows, ncols = len(rows), len(rows[0])
    tbl = doc.add_table(rows=nrows, cols=ncols)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        try:
            tbl.style = "Table Normal"
        except KeyError:
            pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = tbl.rows[r].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                _format_paragraph(
                    p,
                    size=TABLE_SIZE,
                    bold=(r == 0),
                    align=WD_ALIGN_PARAGRAPH.CENTER if c > 0 and r > 0 else WD_ALIGN_PARAGRAPH.LEFT,
                    space_after=0,
                )
            if r == 0:
                _shade_cell(cell, "E7EEF7")

    spacer = doc.add_paragraph()
    _format_paragraph(spacer, space_after=6)
    return [cap._p, tbl._tbl, spacer._p]


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    elems = _make_table_element(doc, caption, rows)
    for el in elems:
        doc.element.body.append(el)


def _scale_inline_images(paragraph, width_in=6.2) -> None:
    target_cx = int(Inches(width_in))
    for inline in paragraph._element.xpath(".//wp:inline"):
        extents = inline.xpath(".//wp:extent")
        if not extents:
            continue
        ext = extents[0]
        cx = int(ext.get("cx", target_cx))
        cy = int(ext.get("cy", target_cx))
        if cx <= 0:
            cx = target_cx
        ratio = cy / cx
        ext.set("cx", str(target_cx))
        ext.set("cy", str(int(target_cx * ratio)))


def polish_document(docx: Path, blind: bool = False) -> None:
    doc = Document(str(docx))

    kill = (
        "Minghao Li",
        "061123071@stu.aqnu.edu.cn",
        "Anqing Normal University",
        "Corresponding author",
        "CRediT authorship contribution statement",
    ) if blind else ()

    # Remove author-identifying paragraphs
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if blind and (any(k in text for k in kill) or text.startswith("Minghao Li:")):
            p._element.getparent().remove(p._element)

    # Remove broken pandoc tables and figures (re-inserted below)
    for tbl in list(doc.tables):
        tbl._element.getparent().remove(tbl._element)
    _strip_pandoc_figures(doc)

    # Remove mangled auto table captions from pandoc
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if p.style.name == "Table Caption":
            p._element.getparent().remove(p._element)
        elif t.startswith("Table 1.") or t.startswith("Table 2.") or t.startswith("Table 3.") or t.startswith("Table 4.") or t.startswith("Table A"):
            p._element.getparent().remove(p._element)

    # Typography pass + keywords fix
    for p in doc.paragraphs:
        text = _replace_unicode_dashes(p.text.strip())

        if text.startswith("agricultural UAV") and "Introduction" not in text and p.style.name != "Heading 1":
            p.text = (
                "Keywords: agricultural UAV; image restoration; semantic segmentation; "
                "precision agriculture; task-oriented evaluation; crop/weed mapping"
            )

        style = p.style.name
        if style == "Title":
            _format_paragraph(p, size=TITLE_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        elif style == "Abstract Title":
            p.text = "Abstract"
            _format_paragraph(p, size=H2_SIZE, bold=True, space_before=6, space_after=6)
        elif style == "Abstract":
            _format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        elif style.startswith("Heading"):
            pass  # numbered later in _number_headings
        elif style in ("Image Caption", "Caption"):
            _format_paragraph(p, size=CAPTION_SIZE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif style == "Bibliography":
            _format_paragraph(p, size=Pt(11), space_after=3)
        else:
            _format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        _scale_inline_images(p)

    # Highlights belong on Title page / EM field only — remove from manuscript body.
    hl_markers = (
        "Task utility",
        "Leakage-free frozen SegFormer",
        "mIoU gains on WeedsGalore",
        "Strong SR baselines cut BRISQUE",
        "Lightweight 25.7K restorer",
    )
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t == "Highlights" or any(t.lstrip("0123456789. ").startswith(m) for m in hl_markers):
            p._element.getparent().remove(p._element)

    _insert_content_blocks(doc)
    _insert_table_after_caption(doc, "Table A3.", 5)
    _number_headings(doc)

    doc.save(str(docx))


def rebuild_blind_manuscript() -> None:
    """High-quality blind Word: pandoc body + reinserted tables + typography pass."""
    build_reference_doc()
    blind = write_blind_tex()
    pandoc_tex = write_pandoc_tex(blind)
    raw = OUT / "_manuscript_blind_raw.docx"
    pandoc_to_docx(pandoc_tex, raw)
    out = OUT / "manuscript_blind_sat.docx"
    shutil.copy2(raw, out)
    polish_document(out, blind=True)


def rebuild_full_manuscript() -> None:
    build_reference_doc()
    raw = OUT / "_manuscript_raw.docx"
    pandoc_to_docx(PAPER / "main.tex", raw)
    out = OUT / "manuscript_sat.docx"
    shutil.copy2(raw, out)
    polish_document(out, blind=False)


def build_title_page() -> None:
    doc = _new_styled_document()
    t = doc.add_paragraph(MANUSCRIPT_TITLE)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_paragraph(t, size=TITLE_SIZE, bold=True, space_after=12)

    for line in [
        "Minghao Li",
        "Anqing Normal University, Anqing, China",
        "Corresponding author: Minghao Li",
        "Email: 061123071@stu.aqnu.edu.cn",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _format_paragraph(p, space_after=3)

    doc.add_paragraph()
    h = doc.add_paragraph("Abstract")
    _format_paragraph(h, size=H2_SIZE, bold=True)
    ap = doc.add_paragraph(_read_abstract_text())
    _format_paragraph(ap, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    h = doc.add_paragraph("Keywords")
    _format_paragraph(h, size=H2_SIZE, bold=True)
    kp = doc.add_paragraph(
        "agricultural UAV; image restoration; semantic segmentation; precision agriculture; "
        "task-oriented evaluation; crop/weed mapping"
    )
    _format_paragraph(kp)

    h = doc.add_paragraph("Highlights")
    _format_paragraph(h, size=H2_SIZE, bold=True)
    for line in (PAPER / "highlights_sat.txt").read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and line[0].isdigit() and "." in line[:3]:
            bp = doc.add_paragraph(line.split(".", 1)[1].strip())
            _format_paragraph(bp)

    doc.save(str(OUT / "title_page_sat.docx"))


def build_humans_animals_statement() -> None:
    doc = _new_styled_document()
    h = doc.add_paragraph("Statement for Studies in Humans and Animals")
    _format_paragraph(h, size=H1_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    for line in [
        f"Manuscript title: {MANUSCRIPT_TITLE}",
        "Author: Minghao Li, Anqing Normal University, Anqing, China",
        "Date: 22 June 2026",
        "",
        "This manuscript does not report studies involving human participants, human data, or "
        "animal subjects. The work uses agricultural UAV imagery and crop/weed image datasets for "
        "image restoration and perception experiments.",
    ]:
        p = doc.add_paragraph(line)
        _format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY if line.startswith("This manuscript") else None)
    for target in (PAPER / "statement_humans_animals_sat.docx", OUT / "statement_humans_animals_sat.docx"):
        doc.save(str(target))


def build_tables_docx() -> None:
    doc = _new_styled_document()
    h = doc.add_paragraph("Tables for typesetting")
    _format_paragraph(h, size=H1_SIZE, bold=True)
    intro = doc.add_paragraph(
        "Editable table source for Smart Agricultural Technology (matches manuscript numbering)."
    )
    _format_paragraph(intro)
    for caption, rows in MANUSCRIPT_TABLES[:4]:
        add_table(doc, caption, rows)
    doc.save(str(OUT / "tables_sat.docx"))


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_reference_doc()
    rebuild_blind_manuscript()
    rebuild_full_manuscript()
    build_title_page()
    build_tables_docx()
    build_humans_animals_statement()

    for name in [
        "cover_letter_sat.docx",
        "cover_letter_sat.pdf",
        "declaration_competing_interests_sat.docx",
        "abstract_sat.docx",
        "highlights_sat.docx",
    ]:
        src = PAPER / name
        if src.exists():
            shutil.copy2(src, OUT / name)

    print("Wrote polished revision package to", OUT)
    for p in sorted(OUT.iterdir()):
        if p.name.startswith("_") or p.name.startswith(".~") or p.name == "test2.docx":
            continue
        print(f"  {p.name}\t{p.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
